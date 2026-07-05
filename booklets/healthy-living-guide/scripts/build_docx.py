#!/usr/bin/env python3
"""Build the editable Word (.docx) version of the Healthy Living Guide."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from docx_helpers import (
    heading, body, kicker, add_callout, add_checklist, add_plain_bullets,
    add_image_centered, add_data_table, DEEP_BLUE, MID_BLUE, TEAL, GREEN, SLATE, INK
)

ROOT = os.path.join(os.path.dirname(__file__), "..")
GFX = os.path.join(ROOT, "graphics", "png", "infographics")
ICONS = os.path.join(ROOT, "graphics", "png", "icons")

doc = Document()

# ---- base style ----
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = INK

for sec in doc.sections:
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

# ============================================================ COVER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NEBRASKA MEDICINE HEALTH SERIES · BOOK ONE")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = TEAL

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Healthy Living Guide")
run.bold = True
run.font.size = Pt(40)
run.font.color.rgb = DEEP_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Your Everyday Blueprint for Living Longer and Feeling Better")
run.font.size = Pt(15)
run.italic = True
run.font.color.rgb = SLATE

doc.add_paragraph()
auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = auth.add_run("Dr. Jason Patera")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = DEEP_BLUE
auth2 = doc.add_paragraph()
auth2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = auth2.add_run("Family Medicine · Nebraska Medicine")
run.font.size = Pt(11)
run.font.color.rgb = SLATE

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Based on guidance from USPSTF, AHA/ACC, ADA, ACSM, the American Cancer Society, "
                    "CDC, and the Dietary Guidelines for Americans.")
run.font.size = Pt(9.5)
run.italic = True
run.font.color.rgb = SLATE

doc.add_paragraph()
kicker(doc, "What's Inside", color=DEEP_BLUE)
toc_items = [
    "The 10 Habits That Matter Most", "Eat Mostly Real Food", "Build Better Meals",
    "Healthy Snacks", "Exercise", "Healthy Weight", "Sleep, Stress & Connection",
    "Preventive Health", "Healthy Grocery Guide", "Frequently Asked Questions",
]
for i, t in enumerate(toc_items, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(t)

doc.add_page_break()

# ============================================================ SECTION 1
kicker(doc, "Start Here", color=TEAL)
heading(doc, "The 10 Habits That Matter Most", level=1)
body(doc, "You don't need to change everything at once. These ten habits have the strongest "
          "evidence for helping you live longer and feel better — pick one to start tomorrow.",
     size=12, color=SLATE)

habits = [
    ("1. Move every day", "150–300 minutes a week of moderate activity, plus strength training twice a week, lowers your risk of heart disease, diabetes, and early death.", "Take a 10-minute walk after dinner."),
    ("2. Eat mostly real food", "Build meals around vegetables, fruits, whole grains, lean protein, and healthy fats instead of packaged, ultra-processed foods.", "Choose foods with a short ingredient list."),
    ("3. Fill half your plate with produce", "Vegetables and fruit provide fiber, vitamins, and antioxidants that protect your heart, gut, and metabolism.", "Add a vegetable to two meals today."),
    ("4. Make protein a priority", "Protein helps preserve muscle, control appetite, and support healthy aging — at every stage of life.", "Include a protein source at every meal."),
    ("5. Choose water first", "Water supports nearly every process in your body and has zero calories or added sugar.", "Keep a water bottle where you'll see it."),
    ("6. Protect 7–9 hours of sleep", "Sleep restores your brain, balances hunger hormones, and supports a healthy weight and mood.", "Set a consistent bedtime, even on weekends."),
    ("7. Build strength & balance", "Strength training preserves muscle and bone as you age and lowers your risk of falls and injury.", "Start with two short sessions this week."),
    ("8. Nurture your mind & relationships", "Managing stress and staying socially connected are linked to a longer, healthier life.", "Reach out to one person you care about today."),
    ("9. Know your numbers", "Blood pressure, cholesterol, blood sugar, and weight trends are early warning signs your body gives you.", "Ask your care team what your numbers are."),
    ("10. Stay current on screenings & vaccines", "Preventive care catches problems early — when they are most treatable.", "See the Preventive Health section for the schedule that fits your age."),
]
for name, desc, tip in habits:
    h = doc.add_paragraph()
    r = h.add_run(name)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = DEEP_BLUE
    body(doc, desc, size=10.5)
    p = doc.add_paragraph()
    r = p.add_run("Try this: ")
    r.bold = True
    r.font.color.rgb = TEAL
    r2 = p.add_run(tip)
    r2.font.color.rgb = TEAL
    p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# ============================================================ SECTION 2
kicker(doc, "Nutrition", color=TEAL)
heading(doc, "Eat Mostly Real Food", level=1)
body(doc, "There is no single perfect diet — but a Mediterranean-style pattern, built mostly "
          "from whole foods, has the strongest evidence for a longer, healthier life.", size=12, color=SLATE)
add_image_centered(doc, os.path.join(GFX, "healthy_plate.png"), 4.2)

add_data_table(doc, ["Nutrient", "Choose more often", "Enjoy less often"], [
    ["Protein", "Fish, poultry, eggs, beans, lentils, tofu, Greek yogurt", "Bacon, sausage, deli & cured meats"],
    ["Fiber", "Vegetables, fruit, beans, oats, whole grains", "Refined white bread & pastries"],
    ["Fats", "Olive oil, nuts, seeds, avocado, fatty fish", "Fried foods, stick margarine"],
    ["Carbohydrates", "100% whole grains, legumes, starchy vegetables", "Sugary cereal, refined snack foods"],
    ["Added sugar", "Whole fruit for something sweet", "Soda, candy, sweetened coffee drinks"],
    ["Processed foods", "Home-prepared, minimally processed meals", "Packaged snacks, fast food, sugary drinks"],
])
doc.add_paragraph()
add_callout(doc, "Did you know?", "Only about 5% of American adults eat enough fiber. Most need 25–38 grams a day.", "tip")
add_callout(doc, "Did you know?", "Large clinical trials show a Mediterranean eating pattern lowers the risk of heart attack and stroke.", "tip")
add_callout(doc, "What to do tomorrow", "No food group is off-limits. Progress comes from what you add — more vegetables, more fiber, more protein — not from what you eliminate.", "action")

doc.add_page_break()

# ============================================================ SECTION 3
kicker(doc, "Nutrition", color=TEAL)
heading(doc, "Build Better Meals", level=1)
body(doc, "A simple formula works for almost any meal: protein + fiber + healthy fat + color.", size=12, color=SLATE)

heading(doc, "Breakfast", level=2, size=13)
add_checklist(doc, ["Greek yogurt, berries & walnuts", "Eggs, whole-grain toast & avocado",
                     "Oatmeal, peanut butter & banana", "Cottage cheese & fruit"])
heading(doc, "Lunch & Dinner", level=2, size=13)
add_checklist(doc, ["Palm-sized protein", "Half the plate vegetables",
                     "Quarter plate whole grain or starchy veg", "A drizzle of olive oil"])
heading(doc, "Meal Planning", level=2, size=13)
add_checklist(doc, ["Plan 3–4 dinners before you shop", "Cook once, eat twice — make extra",
                     "Keep a running grocery list", "Prep produce right after shopping"])

add_image_centered(doc, os.path.join(GFX, "nutrition_label.png"), 3.6)

heading(doc, "Eating Out", level=2, size=13)
add_plain_bullets(doc, ["Look at the menu ahead of time and decide before you arrive",
                         "Ask for dressing and sauce on the side",
                         "Choose grilled, baked, or roasted over fried",
                         "Start with a vegetable-based side or salad",
                         "Box up half your meal before you start eating",
                         "Choose water or unsweetened tea over soda"])

doc.add_page_break()

# ============================================================ SECTION 4
kicker(doc, "Nutrition", color=TEAL)
heading(doc, "Healthy Snacks", level=1)
body(doc, "A good snack pairs protein or fiber with something you enjoy. Values are approximate "
          "and will vary by brand and portion.", size=12, color=SLATE)

snacks = [
    ("Apple + 1 tbsp peanut butter", "P 4g · F 4g · ~190 Cal", "Fiber + protein keeps you full"),
    ("Plain Greek yogurt (3/4 cup)", "P 17g · F 0g · ~100 Cal", "High protein, gut-healthy"),
    ("Almonds (1 oz, ~23)", "P 6g · F 3.5g · ~160 Cal", "Healthy fats + fiber"),
    ("Baby carrots + hummus (2 tbsp)", "P 3g · F 4g · ~100 Cal", "Crunchy fiber boost"),
    ("Hard-boiled egg (1)", "P 6g · F 0g · ~70 Cal", "Portable protein"),
    ("Cottage cheese + pineapple", "P 12g · F 1g · ~120 Cal", "Slow-digesting protein"),
    ("Edamame, steamed (1 cup)", "P 17g · F 8g · ~190 Cal", "Plant protein + fiber"),
    ("String cheese (1 stick)", "P 7g · F 0g · ~80 Cal", "Quick, portable protein"),
    ("Trail mix, unsweetened (1/4 cup)", "P 5g · F 2.5g · ~170 Cal", "Energy + healthy fats"),
    ("Banana + 1 tbsp almond butter", "P 4g · F 4g · ~200 Cal", "Potassium + protein"),
    ("Whole-grain crackers + cheese", "P 6g · F 2g · ~150 Cal", "Balanced carb + protein"),
    ("Roasted chickpeas (1/2 cup)", "P 6g · F 5g · ~120 Cal", "Crunchy plant fiber"),
    ("Popcorn, air-popped (3 cups)", "P 3g · F 3.5g · ~90 Cal", "Whole grain, low calorie"),
    ("Cherry tomatoes + mozzarella", "P 7g · F 1g · ~110 Cal", "Protein + produce"),
    ("Protein smoothie (small)", "P 20g · F 2g · ~150 Cal", "Convenient post-activity fuel"),
    ("Orange + 12 almonds", "P 4g · F 5g · ~150 Cal", "Vitamin C + healthy fat"),
    ("Tuna (2 oz) + whole-grain crackers", "P 12g · F 2g · ~140 Cal", "Omega-3s + protein"),
    ("Cucumber slices + tzatziki", "P 3g · F 2g · ~90 Cal", "Low-cal hydration + fiber"),
    ("Overnight oats (1/2 cup)", "P 6g · F 4g · ~150 Cal", "Fiber-rich, make ahead"),
    ("Dark chocolate square + walnuts", "P 3g · F 2g · ~140 Cal", "Satisfying, antioxidant-rich"),
]
add_data_table(doc, ["Snack", "Protein / Fiber / Calories", "Why it helps"],
               [[a, b, c] for a, b, c in snacks])
doc.add_paragraph()
add_callout(doc, "What to do tomorrow", "Keep two or three of these on hand at home, at work, and in your bag so a healthy option is always within reach.", "action")

doc.add_page_break()

# ============================================================ SECTION 5
kicker(doc, "Movement", color=TEAL)
heading(doc, "Exercise", level=1)
body(doc, "Adults need 150–300 minutes of moderate aerobic activity a week, plus strength "
          "training at least twice a week. Any amount of movement is better than none.", size=12, color=SLATE)
add_image_centered(doc, os.path.join(GFX, "exercise_pyramid.png"), 4.0)
add_image_centered(doc, os.path.join(GFX, "zone2_chart.png"), 5.5)
add_image_centered(doc, os.path.join(GFX, "weekly_calendar.png"), 6.0)

heading(doc, "How to begin, at any fitness level", level=2, size=13)
add_checklist(doc, ["Start with 10 minutes of walking, most days", "Add 2 short strength sessions a week",
                     "Use the talk test to find your Zone 2 pace", "Add 5 minutes of mobility work after warming up",
                     "Increase time or distance by about 10% a week", "Choose an activity you actually enjoy — it will last longer"])

doc.add_page_break()

# ============================================================ SECTION 6
kicker(doc, "Movement & Nutrition", color=TEAL)
heading(doc, "Healthy Weight", level=1)
add_callout(doc, "Did you know?", "Losing just 5–10% of body weight can improve blood pressure, blood sugar, fatty liver disease, and joint pain.", "tip")

heading(doc, "Timing & Fasting", level=2, size=13)
body(doc, "What you eat matters more than exactly when you eat. Intermittent fasting can be a "
          "helpful tool for some people to manage calorie intake, but it isn't required and "
          "isn't right for everyone — including those who are pregnant, have a history of an "
          "eating disorder, or take medications for diabetes. Ask your care team first.", size=10.5)

heading(doc, "GLP-1 Medications", level=2, size=13)
body(doc, "GLP-1 medications can be an effective option for some adults with obesity or type 2 "
          "diabetes. They work best alongside — not instead of — healthy eating, strength "
          "training, and sleep, which help preserve muscle while losing weight. Talk with your "
          "care team about whether one is right for you.", size=10.5)

add_image_centered(doc, os.path.join(GFX, "weight_habits_wheel.png"), 4.2)

heading(doc, "Behavior Change That Lasts", level=2, size=13)
add_checklist(doc, ["Change one habit at a time", "Track progress with more than the scale — energy, sleep, and clothes fit all count",
                     "Protect your sleep — poor sleep increases hunger hormones", "Manage stress — it drives cravings and emotional eating"])
add_callout(doc, "Protect your muscle", "Whether or not you use medication, pairing weight loss with adequate protein and regular strength training helps preserve muscle and keep your metabolism strong.", "action")

doc.add_page_break()

# ============================================================ SECTION 7
kicker(doc, "Well-Being", color=TEAL)
heading(doc, "Sleep, Stress & Connection", level=1)
body(doc, "Most adults need 7–9 hours of sleep a night. Sleep, stress, and relationships all "
          "shape your physical health.", size=12, color=SLATE)

heading(doc, "Sleep Hygiene", level=2, size=13)
add_checklist(doc, ["Keep a consistent sleep & wake time", "Keep your room cool, dark & quiet",
                     "Limit screens for 30+ minutes before bed", "Avoid caffeine after early afternoon",
                     "Get morning sunlight when you can"])
heading(doc, "Managing Stress", level=2, size=13)
add_checklist(doc, ["Try slow breathing for 2 minutes", "Take a short walk outside",
                     "Write down what's on your mind", "Limit news & social media before bed",
                     "Notice tension without judging it"])
heading(doc, "Social Connection", level=2, size=13)
add_checklist(doc, ["Reach out to one person most days", "Say yes to shared meals & activities",
                     "Join a group tied to an interest", "Ask for help — and offer it"])

add_callout(doc, "Did you know?", "The U.S. Surgeon General has identified loneliness and social isolation as risk factors for heart disease, stroke, and dementia — comparable in impact to other well-known health risks.", "tip")
add_callout(doc, "When to reach out for help", "If low mood, anxiety, or stress are affecting your daily life for more than two weeks, talk with your care team. Effective treatment is available, and asking for help is a sign of strength, not weakness.", "action")

heading(doc, "A 2-Minute Reset for a Stressful Moment", level=2, size=13)
add_plain_bullets(doc, ["1. Pause — stop what you're doing and sit or stand comfortably.",
                         "2. Breathe — inhale for 4 counts, hold for 4, exhale slowly for 6. Repeat 5 times.",
                         "3. Reset — name one thing you can do next, then take that one small step."])

doc.add_page_break()

# ============================================================ SECTION 8
kicker(doc, "Prevention", color=TEAL)
heading(doc, "Preventive Health", level=1)
body(doc, "Screenings and vaccines catch problems early. Ask your care team which of these apply to you.", size=12, color=SLATE)

add_data_table(doc, ["Screening", "Who", "How often"], [
    ["Blood pressure", "All adults", "Every visit; at least annually"],
    ["Cholesterol", "Adults 20+", "Every 4–6 years, more often with risk factors"],
    ["Blood sugar / diabetes", "Adults 35–70 who are overweight, or with other risk factors", "Every 3 years"],
    ["Breast cancer (mammogram)", "Women 40–74", "Every 2 years"],
    ["Cervical cancer", "Women 21–65", "Every 3–5 years (Pap/HPV)"],
    ["Colorectal cancer", "Adults 45–75", "Colonoscopy every 10 yrs, or annual stool test"],
    ["Lung cancer (low-dose CT)", "Adults 50–80 with a significant smoking history", "Annually"],
    ["Dental exam & cleaning", "All adults", "Every 6–12 months"],
    ["Vision exam", "All adults", "Every 1–2 years, more often after age 60"],
    ["Hearing check", "All adults", "Baseline, then as needed"],
    ["Skin check", "All adults, especially with risk factors", "Annually, or if a mole changes"],
])
doc.add_paragraph()
heading(doc, "Vaccines", level=2, size=13)
add_plain_bullets(doc, ["Influenza — every year", "Tdap/Td — once, then booster every 10 years",
                         "Shingles — age 50+", "Pneumococcal — age 65+, or earlier with certain conditions",
                         "COVID-19 — per current CDC guidance"])
add_callout(doc, "Sun & skin protection", "Use SPF 30+ daily, seek shade midday, and check your skin monthly for new or changing spots — remember the ABCDEs: Asymmetry, Border, Color, Diameter, Evolving.", "action")

doc.add_page_break()

# ============================================================ SECTION 9
kicker(doc, "Nutrition", color=TEAL)
heading(doc, "Healthy Grocery Guide", level=1)
add_image_centered(doc, os.path.join(GFX, "grocery_cart.png"), 4.3)
add_image_centered(doc, os.path.join(GFX, "healthy_pantry.png"), 5.5)

heading(doc, "Simple Swaps", level=2, size=13)
add_data_table(doc, ["Instead of", "Choose more often"], [
    ["White bread", "100% whole-grain or whole-wheat bread"],
    ["Sugary breakfast cereal", "Oats or a lower-sugar, whole-grain cereal"],
    ["Soda or sweetened tea", "Sparkling water or unsweetened tea"],
    ["Potato chips", "Air-popped popcorn or nuts"],
    ["Regular pasta", "Whole-wheat or legume-based pasta"],
    ["Creamy salad dressing", "Olive oil & vinegar"],
    ["Frozen meals, highly processed", "Plain frozen vegetables & proteins you season yourself"],
])

doc.add_page_break()

# ============================================================ SECTION 10
kicker(doc, "Your Questions", color=TEAL)
heading(doc, "Frequently Asked Questions", level=1)

faqs = [
    ("Is fruit too sugary?", "No. The fiber in whole fruit slows sugar absorption. Whole fruit is linked to better, not worse, health outcomes."),
    ("Are eggs bad for cholesterol?", "For most people, eggs in moderation (about one a day) fit well into a healthy eating pattern."),
    ("Is coffee healthy?", "Moderate coffee (3–4 cups/day) is linked to neutral or lower risk of several conditions. Watch the added sugar and cream."),
    ("Are artificial sweeteners safe?", "FDA-approved sweeteners are generally safe in moderation and can help reduce added sugar. Water is still the best default."),
    ("Do I need protein shakes?", "Not necessarily. They're a convenient option if you struggle to get enough protein from food — choose a lower-sugar option."),
    ("How much alcohol is safe?", "Less is better. If you drink, current guidance suggests limiting to 2 drinks/day or fewer for men and 1 or fewer for women."),
    ("Do I have to eat breakfast?", "Not necessarily — it depends on the person. If you do eat breakfast, prioritize protein and fiber to help you feel full."),
    ("Do I need a multivitamin?", "Most nutrients are best absorbed from food. Some people benefit from specific supplements (like vitamin D) — ask your care team."),
    ("Does intermittent fasting work?", "It can help some people manage calorie intake, but it isn't more effective than other approaches when calories are matched, and isn't right for everyone."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DEEP_BLUE
    body(doc, a, size=10.5)

add_callout(doc, "Still have questions?", "Bring this guide to your next visit. Your care team can help you turn any of these habits into a plan that fits your life.", "action")

doc.add_page_break()

# ============================================================ REFERENCES
kicker(doc, "Evidence Basis", color=TEAL)
heading(doc, "References & Guidelines", level=1)
body(doc, "This guide reflects current guidance from major medical and public health "
          "organizations. It is educational and does not replace personalized medical advice "
          "from your care team.", size=11, color=SLATE)

ref_groups = [
    ("Nutrition", ["U.S. Dietary Guidelines for Americans, 2020–2025", "American Heart Association dietary guidance", "PREDIMED & related Mediterranean diet trials"]),
    ("Exercise", ["American College of Sports Medicine (ACSM) guidelines", "Physical Activity Guidelines for Americans, HHS"]),
    ("Weight & Diabetes", ["American Diabetes Association (ADA) Standards of Care", "AHA/ACC/Obesity Society guidance"]),
    ("Cancer Screening", ["American Cancer Society screening guidelines", "U.S. Preventive Services Task Force (USPSTF)"]),
    ("Cardiovascular Health", ["American Heart Association / American College of Cardiology (AHA/ACC)"]),
    ("Sleep & Mental Health", ["American Academy of Sleep Medicine", "U.S. Surgeon General's Advisory on Loneliness (2023)"]),
    ("General Prevention", ["Centers for Disease Control and Prevention (CDC)", "U.S. Preventive Services Task Force (USPSTF)"]),
]
for title, items in ref_groups:
    h = doc.add_paragraph()
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = TEAL
    add_plain_bullets(doc, items)

doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run("About the Author")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = DEEP_BLUE
body(doc, "Dr. Jason Patera is a family medicine physician with Nebraska Medicine. This guide "
          "is part of a series of patient education booklets designed to help patients build "
          "sustainable, evidence-based habits for lifelong health.", size=10.5)

doc.add_paragraph()
body(doc, "© Nebraska Medicine. This material is for general educational purposes and is not a "
          "substitute for professional medical advice, diagnosis, or treatment. Always ask your "
          "physician or qualified health provider with questions about a medical condition.",
     size=8.5, color=SLATE, italic=True)

OUT = os.path.join(ROOT, "build", "Healthy-Living-Guide.docx")
doc.save(OUT)
print("Wrote", OUT)
