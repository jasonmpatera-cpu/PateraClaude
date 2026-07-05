from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DEEP_BLUE = RGBColor(0x0F, 0x3D, 0x5C)
MID_BLUE = RGBColor(0x1B, 0x6C, 0xA8)
TEAL = RGBColor(0x12, 0x80, 0x74)
GREEN = RGBColor(0x1E, 0x9E, 0x6B)
SLATE = RGBColor(0x4B, 0x5F, 0x68)
INK = RGBColor(0x1C, 0x2B, 0x33)

LIGHT_BLUE_HEX = "E8F2FA"
LIGHT_GREEN_HEX = "E6F6F0"
DEEP_BLUE_HEX = "0F3D5C"


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, color="1B6CA8", sz=18, sides=("left",)):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side in sides:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color)
        borders.append(el)
    tcPr.append(borders)


def no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        borders.append(el)
    tblPr.append(borders)


def add_callout(doc, label, text, kind="tip"):
    """kind: 'tip' (blue) or 'action' (green)"""
    color_hex = LIGHT_BLUE_HEX if kind == "tip" else LIGHT_GREEN_HEX
    border_hex = "1B6CA8" if kind == "tip" else "1E9E6B"
    label_color = MID_BLUE if kind == "tip" else TEAL

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, color_hex)
    set_cell_border(cell, color=border_hex, sz=24, sides=("left",))
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    p_label = cell.paragraphs[0]
    run = p_label.add_run(label.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = label_color

    p_body = cell.add_paragraph()
    run2 = p_body.add_run(text)
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = INK
    p_body.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("✓  ")
        run.bold = True
        run.font.color.rgb = GREEN
        run2 = p.add_run(item)
        run2.font.color.rgb = INK


def add_plain_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_image_centered(doc, path, width_in):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))


def add_data_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(hdr_cells[i], DEEP_BLUE_HEX)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def heading(doc, text, level=1, color=None, size=None):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color or DEEP_BLUE
    if size:
        run.font.size = Pt(size)
    return h


def body(doc, text, size=11, color=None, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or INK
    run.italic = italic
    run.bold = bold
    return p


def kicker(doc, text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = color or TEAL
    return p
